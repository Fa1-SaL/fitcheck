import os
import re
import sys
import hashlib
import zipfile
import xml.etree.ElementTree as ET
from html.parser import HTMLParser
from pathlib import Path
from utils.helpers import logger
from config.config import CACHE_DIR

# Import pdf and docx libraries
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

def clean_text(text: str) -> str:
    """Normalize whitespace and clean up special characters from parsed text."""
    if not text:
        return ""
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n", "\n\n", text)
    return text.strip()

def get_file_hash(file_path: Path) -> str:
    """Compute the SHA-256 hash of a file's binary contents."""
    return hashlib.sha256(file_path.read_bytes()).hexdigest()

def parse_pdf(file_path: Path) -> str:
    """Extract plain text from a PDF file."""
    if not pypdf:
        raise ImportError("The 'pypdf' library is not installed.")
        
    logger.info(f"Parsing PDF file: {file_path.name}")
    reader = pypdf.PdfReader(file_path)
    text_content = []
    
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            text_content.append(page_text)
            
    if not text_content:
        raise ValueError("PDF file returned no text content. The file might be scanned/image-only.")
        
    return clean_text("\n".join(text_content))

def parse_docx(file_path: Path) -> str:
    """Extract plain text from a DOCX file."""
    if not docx:
        raise ImportError("The 'python-docx' library is not installed.")
        
    logger.info(f"Parsing DOCX file: {file_path.name}")
    doc = docx.Document(str(file_path))
    text_content = []
    
    # Extract from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_content.append(para.text)
            
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_content.append(" | ".join(row_text))
                
    if not text_content:
        raise ValueError("DOCX file contains no readable text paragraphs or tables.")
        
    return clean_text("\n".join(text_content))

class HTMLTextExtractor(HTMLParser):
    def __init__(self):
        super().__init__()
        self.texts = []
        self.ignore = False
    def handle_starttag(self, tag, attrs):
        if tag in ["script", "style", "head", "noscript"]:
            self.ignore = True
    def handle_endtag(self, tag):
        if tag in ["script", "style", "head", "noscript"]:
            self.ignore = False
        if tag in ["p", "div", "br", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6"]:
            self.texts.append("\n")
    def handle_data(self, data):
        if not self.ignore and data.strip():
            self.texts.append(data.strip())

def parse_html(file_path: Path) -> str:
    """Extract plain text from an HTML/HTM file."""
    logger.info(f"Parsing HTML file: {file_path.name}")
    try:
        content = file_path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        content = file_path.read_bytes().decode("latin-1", errors="ignore")
    extractor = HTMLTextExtractor()
    extractor.feed(content)
    cleaned = clean_text(" ".join(extractor.texts))
    if not cleaned:
        raise ValueError("HTML file yielded no readable text.")
    return cleaned

def parse_rtf(file_path: Path) -> str:
    """Extract plain text from an RTF file."""
    logger.info(f"Parsing RTF file: {file_path.name}")
    try:
        content = file_path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        content = file_path.read_bytes().decode("latin-1", errors="ignore")
    text = re.sub(r'\\(par|line|row)\b', '\n', content)
    text = re.sub(r'\\[a-z]+(-?\d+)?( ?)', '', text)
    text = re.sub(r'\{[^{}]*\}', '', text)
    text = re.sub(r'[{}]', '', text)
    cleaned = clean_text(text)
    if not cleaned:
        raise ValueError("RTF file yielded no readable text.")
    return cleaned

def parse_odt(file_path: Path) -> str:
    """Extract plain text from an OpenDocument Text (.odt) file."""
    logger.info(f"Parsing ODT file: {file_path.name}")
    with zipfile.ZipFile(file_path, "r") as zf:
        xml_content = zf.read("content.xml")
    tree = ET.fromstring(xml_content)
    texts = []
    for elem in tree.iter():
        if elem.text and elem.text.strip():
            texts.append(elem.text.strip())
        if elem.tail and elem.tail.strip():
            texts.append(elem.tail.strip())
    cleaned = clean_text("\n".join(texts))
    if not cleaned:
        raise ValueError("ODT file yielded no readable text.")
    return cleaned

def parse_txt(file_path: Path) -> str:
    """Extract plain text from TXT or raw text files."""
    logger.info(f"Parsing TXT file: {file_path.name}")
    content = file_path.read_bytes()
    for enc in ["utf-8", "utf-16", "latin-1", "cp1252"]:
        try:
            text = content.decode(enc)
            cleaned = clean_text(text)
            if cleaned:
                return cleaned
        except Exception:
            continue
    cleaned = clean_text(content.decode("utf-8", errors="ignore"))
    if not cleaned:
        raise ValueError("Text file yielded no readable text.")
    return cleaned

def parse_doc_fallback(file_path: Path) -> str:
    """
    Fallback parser for binary legacy DOC files or unrecognised formats.
    Reads binary content and extracts readable ASCII/UTF-8 string sequences.
    """
    logger.warning(f"Using binary text extraction fallback parser for file: {file_path.name}")
    try:
        content = file_path.read_bytes()
        printables = re.findall(rb"[\x20-\x7E\x09\x0A\x0D]{4,}", content)
        text = b" \n ".join(printables).decode("ascii", errors="ignore")
        text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)
        cleaned = clean_text(text)
        if len(cleaned) < 20:
            raise ValueError("ASCII extraction yielded too little text.")
        return cleaned
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from {file_path.name}: {str(e)}")

def parse_doc_win32(file_path: Path) -> str:
    """Extract plain text from a legacy .doc file using pywin32 and MS Word."""
    logger.info(f"Attempting to parse legacy DOC file via MS Word (pywin32): {file_path.name}")
    try:
        import win32com.client
    except ImportError:
        logger.warning("pywin32 is not installed or not supported on this platform.")
        return parse_doc_fallback(file_path)
        
    word_app = None
    doc = None
    try:
        import pythoncom
        pythoncom.CoInitialize()
        
        word_app = win32com.client.Dispatch("Word.Application")
        word_app.Visible = False
        
        abs_path = str(file_path.resolve())
        doc = word_app.Documents.Open(abs_path, ReadOnly=True)
        text = doc.Content.Text
        
        doc.Close(False)
        word_app.Quit()
        return clean_text(text)
    except Exception as e:
        logger.error(f"MS Word automation failed: {str(e)}")
        try:
            if doc:
                doc.Close(False)
        except:
            pass
        try:
            if word_app:
                word_app.Quit()
        except:
            pass
        return parse_doc_fallback(file_path)

def parse_doc(file_path: Path) -> str:
    """Routes legacy doc parsing based on OS and availability of win32com."""
    if sys.platform == "win32":
        return parse_doc_win32(file_path)
    else:
        return parse_doc_fallback(file_path)

def extract_text_from_file(file_path: Path) -> str:
    """
    Extract plain text from resume file supporting PDF, DOCX, DOC, Google Doc, RTF, ODT, HTML, TXT and more.
    Uses magic bytes and multi-tier fallbacks to ensure successful text extraction.
    Caches parsed text locally using file content hash.
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
        
    # Calculate file content hash
    file_hash = get_file_hash(file_path)
    cache_file = CACHE_DIR / f"{file_hash}_parsed.txt"
    
    # 1. Check Cache
    if cache_file.exists():
        logger.info(f"Reused parsed text for {file_path.name} (found cached parsed text)")
        try:
            return cache_file.read_text(encoding="utf-8")
        except Exception as e:
            logger.warning(f"Failed to read cached text file {cache_file.name}: {str(e)}")
            
    # 2. Inspect Magic Bytes
    try:
        header = file_path.read_bytes()[:2048]
    except Exception:
        header = b""
        
    ext = file_path.suffix.lower()
    header_strip = header.strip()
    
    # Order of parsers to try based on magic bytes or extension
    parsers_to_try = []
    
    if header.startswith(b"%PDF") or ext == ".pdf":
        parsers_to_try.extend([parse_pdf, parse_docx, parse_txt])
    elif header.startswith((b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")) or ext in [".docx", ".odt", ".epub", ".zip"]:
        if ext == ".odt":
            parsers_to_try.extend([parse_odt, parse_docx, parse_txt])
        else:
            parsers_to_try.extend([parse_docx, parse_odt, parse_txt])
    elif header.startswith(b"\xd0\xcf\x11\xe0") or ext == ".doc":
        parsers_to_try.extend([parse_doc, parse_txt, parse_doc_fallback])
    elif header.startswith(b"{\\rtf") or ext == ".rtf":
        parsers_to_try.extend([parse_rtf, parse_txt])
    elif header_strip.startswith((b"<!DOCTYPE html", b"<html", b"<!doctype html")) or ext in [".html", ".htm"]:
        parsers_to_try.extend([parse_html, parse_txt])
    else:
        if ext == ".docx":
            parsers_to_try.extend([parse_docx, parse_txt, parse_doc_fallback])
        elif ext == ".doc":
            parsers_to_try.extend([parse_doc, parse_txt, parse_doc_fallback])
        elif ext == ".pdf":
            parsers_to_try.extend([parse_pdf, parse_txt, parse_doc_fallback])
        elif ext in [".txt", ".csv", ".md"]:
            parsers_to_try.extend([parse_txt, parse_html, parse_doc_fallback])
        else:
            parsers_to_try.extend([parse_txt, parse_html, parse_rtf, parse_docx, parse_doc_fallback])
            
    if parse_doc_fallback not in parsers_to_try:
        parsers_to_try.append(parse_doc_fallback)
        
    errors = []
    text = None
    for parser_func in parsers_to_try:
        try:
            res = parser_func(file_path)
            if res and len(res.strip()) >= 10:
                text = res
                break
        except Exception as e:
            errors.append(f"{parser_func.__name__}: {str(e)}")
            continue
            
    if not text:
        raise RuntimeError(f"Failed to extract text from {file_path.name}. Errors encountered: {'; '.join(errors)}")
        
    # Write parsed text to cache
    try:
        cache_file.write_text(text, encoding="utf-8")
        logger.info(f"Saved parsed text to cache: {cache_file.name}")
    except Exception as e:
        logger.warning(f"Failed to save parsed text cache: {str(e)}")
        
    return text
